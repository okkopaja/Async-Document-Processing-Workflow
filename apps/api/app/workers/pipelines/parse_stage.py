"""Document parsing stage — extract raw text from files."""
from __future__ import annotations

import csv
from dataclasses import dataclass
from pathlib import Path

from app.core.constants import PermanentParseError, TransientParseError, UnsupportedFileTypeError


@dataclass
class ParsedDocument:
	"""Result of parsing a document file."""

	raw_text: str
	line_count: int
	char_count: int
	file_type: str


def parse_stage(file_path: str, mime_type: str, extension: str) -> ParsedDocument:
	"""
	Parse file and extract raw text content.

	Args:
	    file_path: Full path to stored file
	    mime_type: MIME type of file
	    extension: File extension (txt, md, pdf, docx, csv)

	Returns:
	    ParsedDocument with extracted text

	Raises:
	    UnsupportedFileTypeError: If file type is not supported
	    TransientParseError: Temporary parsing error (e.g., I/O error)
	    PermanentParseError: Permanent parsing error (corrupted file)
	"""
	path = Path(file_path)

	if not path.exists():
		raise TransientParseError(f"File not found: {file_path}")

	try:
		extension_lower = extension.lower().lstrip(".")

		# Plain text and markdown
		if extension_lower in ("txt", "md"):
			raw_text = _parse_text(path)

		# PDF handling
		elif extension_lower == "pdf":
			raw_text = _parse_pdf(path)

		# DOCX handling
		elif extension_lower == "docx":
			raw_text = _parse_docx(path)

		# CSV handling
		elif extension_lower == "csv":
			raw_text = _parse_csv(path)

		else:
			raise UnsupportedFileTypeError(f"Unsupported file type: {extension}")

	except UnsupportedFileTypeError:
		raise
	except Exception as exc:
		# Classify as transient vs permanent based on error type
		if isinstance(exc, (OSError, IOError)):
			raise TransientParseError(f"I/O error parsing file: {exc}") from exc
		raise PermanentParseError(f"Error parsing file: {exc}") from exc

	line_count = len(raw_text.splitlines())
	char_count = len(raw_text)

	return ParsedDocument(
		raw_text=raw_text,
		line_count=line_count,
		char_count=char_count,
		file_type=extension_lower,
	)


def _parse_text(path: Path) -> str:
	"""Parse plain text or markdown file."""
	return path.read_text(encoding="utf-8", errors="replace")


def _parse_pdf(path: Path) -> str:
	"""Parse PDF file and extract text.

	Uses pdfplumber if available, otherwise returns mock data.
	"""
	try:
		import pdfplumber
	except ImportError:
		# Mock implementation for development
		return f"[Mock PDF content from {path.name}]\n\nThis is simulated PDF text extraction."

	try:
		raw_text = ""
		with pdfplumber.open(path) as pdf:
			for page in pdf.pages:
				text = page.extract_text()
				if text:
					raw_text += text + "\n"
		return raw_text.strip() or f"[Empty PDF: {path.name}]"
	except Exception as exc:
		raise PermanentParseError(f"PDF parsing failed: {exc}") from exc


def _parse_docx(path: Path) -> str:
	"""Parse DOCX file and extract text.

	Uses python-docx if available, otherwise returns mock data.
	"""
	try:
		from docx import Document
	except ImportError:
		# Mock implementation for development
		return f"[Mock DOCX content from {path.name}]\n\nThis is simulated DOCX text extraction."

	try:
		doc = Document(path)
		paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
		raw_text = "\n".join(paragraphs)
		return raw_text.strip() or f"[Empty DOCX: {path.name}]"
	except Exception as exc:
		raise PermanentParseError(f"DOCX parsing failed: {exc}") from exc


def _parse_csv(path: Path) -> str:
	"""Parse CSV file and return structured representation."""
	try:
		lines = []
		with open(path, "r", encoding="utf-8", errors="replace") as f:
			reader = csv.reader(f)
			# Read first 100 rows for preview
			for i, row in enumerate(reader):
				if i >= 100:
					lines.append("[... truncated ...]")
					break
				if row:
					lines.append(" | ".join(row))

		return "\n".join(lines) or f"[Empty CSV: {path.name}]"
	except Exception as exc:
		raise PermanentParseError(f"CSV parsing failed: {exc}") from exc
